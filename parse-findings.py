#!/usr/bin/env python
'''
Use OWASP Testing Guide V4 spreadsheet, Findings worksheet
to convert to Word document.

If additional filename is passed as argument, uses as an evidence document. Searches evidence to see if supporting evidence exists, based on the presence of test ID (e.g. OTG-INFO-002).
If found, adds a placeholder (since I could not get merge to work). If your evidence is order, it is an easy manual merge (but, yet, manual).

Note that this comes out a bit ugly, but is [arguably] better than 91 tables...
'''

import csv
import re
import sys

from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook

document = Document()

# otg = 0
document.add_heading('Appendix A: Findings', level=1)
document.add_heading('Findings', level=2)

# Load workbook and select Findings sheet

if(len(sys.argv) == 1):
    print('Usage: 'parse-findings.py [Findings.xlsx] [Evidence.docx]')

wb = load_workbook(filename = sys.argv[1])
ws = wb.get_sheet_by_name('Findings')

# Convert rows to findings sections
for row in ws.iter_rows():
    fnd = [cell.value for cell in row]

    document.add_heading(fnd[0] + ' - ' + fnd[5], level=3)
    document.add_heading('Risk: ' + fnd[12], level=4)

    document.add_heading('Finding', level=4)
    document.add_paragraph(fnd[6])

    # Security Threat
    if(fnd[9] > '' and fnd[12] != 'n/a'):
        document.add_heading('Security Threat', level=4)
        document.add_paragraph(fnd[9])

    # Remediation
    if(fnd[11] > ''):
        document.add_heading('Remediation', level=4)
        document.add_paragraph(fnd[11])

    # Testing Technique 
    document.add_heading('Tools/Techniques', level=4)
    document.add_paragraph(fnd[10])

    # Search Evidence document
    if(len(sys.argv) > 2):
        f = open(sys.argv[2], 'rb')
        evi = Document(f)

        regexp = re.compile(fnd[0])

        for ele in evi.paragraphs:
            if(ele.text):
                # Found evidence for test; insert placeholder
                if regexp.search(ele.text):
                    document.add_heading('Evidence', level=4)
                    document.add_paragraph('[TODO: Insert evidence]')

        f.close()
        #document.add_paragraph()

    document.add_paragraph()

document.save('findings.docx')
